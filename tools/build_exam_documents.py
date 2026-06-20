#!/usr/bin/env python3
"""Reorder raw chapter quizzes and build the public Markdown/DOCX artifacts."""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


CHAPTER_NAMES = {
    1: "第一章",
    2: "第二章",
    3: "第三章",
    4: "第四章",
    5: "第五章",
    6: "第六章",
    7: "第七章",
}

TYPE_ORDER = ("single", "multi", "judge")
TYPE_LABELS = {
    "single": "一、单选题",
    "multi": "二、多选题",
    "judge": "三、判断题",
}

BLUE = RGBColor(0x1F, 0x4E, 0x79)
GRAY = RGBColor(0x66, 0x66, 0x66)
INK = RGBColor(0x11, 0x11, 0x11)
LIGHT_FILL = "F3F5F8"
WESTERN_FONT = "Times New Roman"
CHINESE_FONT = "Noto Serif SC"


@dataclass
class Question:
    source_block: int
    source_number: int
    kind: str
    stem: str
    options: str
    answer: str
    explanation: str
    number: int = 0


def clean_inline(text: str) -> str:
    text = text.replace("**", "").replace("__", "")
    text = re.sub(r"(?<!\w)_|_(?!\w)", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def parse_question_line(line: str, kind: str) -> tuple[int, str, str]:
    match = re.match(r"^\*\*(\d+)\.\s*(.*)$", line.strip())
    if not match:
        raise ValueError(f"Cannot parse question line: {line}")

    source_number = int(match.group(1))
    content = clean_inline(match.group(2))

    if kind == "judge":
        return source_number, content, ""

    option_match = re.search(r"(?<![A-Za-z])A\.\s*", content)
    if not option_match:
        raise ValueError(f"Question {source_number} has no options: {line}")

    stem = content[: option_match.start()].strip()
    options = content[option_match.start() :].strip()
    options = re.sub(r"\s+(?=[A-D]\.\s*)", "  ", options)
    return source_number, stem, options


def parse_answer_section(text: str) -> dict[int, tuple[str, str]]:
    text = re.sub(r"\*\*【[^】]*答案】\*\*", "", text)
    marker = re.compile(
        r"(?<![\w.])(\d{1,3})\.\s*\*\*([A-D]+|正确|错误)\*\*。"
    )
    matches = list(marker.finditer(text))
    answers: dict[int, tuple[str, str]] = {}

    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        explanation_text = text[match.end() : end].split("\n\n", 1)[0]
        explanation = clean_inline(explanation_text)
        answers[int(match.group(1))] = (match.group(2), explanation)

    return answers


def parse_raw_quiz(path: Path) -> list[Question]:
    text = path.read_text(encoding="utf-8")
    block_heads = list(
        re.finditer(r"^###\s+一、\s*单项选择题.*$", text, flags=re.MULTILINE)
    )
    if not block_heads:
        raise ValueError(f"No quiz blocks found in {path}")

    questions: list[Question] = []
    for block_index, head in enumerate(block_heads, start=1):
        end = (
            block_heads[block_index].start()
            if block_index < len(block_heads)
            else len(text)
        )
        block = text[head.start() : end]
        answer_head = re.search(r"^###\s+.*答案与.*$", block, flags=re.MULTILINE)
        if not answer_head:
            raise ValueError(f"Block {block_index} has no answer section")

        question_part = block[: answer_head.start()]
        answer_part = block[answer_head.end() :].split("\n---", 1)[0]
        answers = parse_answer_section(answer_part)

        current_kind: str | None = None
        block_questions: list[Question] = []
        for line in question_part.splitlines():
            if re.match(r"^###\s+.*单项选择题", line):
                current_kind = "single"
                continue
            if re.match(r"^###\s+.*多项选择题", line):
                current_kind = "multi"
                continue
            if re.match(r"^###\s+.*判断题", line):
                current_kind = "judge"
                continue
            if not re.match(r"^\*\*\d+\.", line):
                continue
            if current_kind is None:
                raise ValueError(f"Question outside a type section: {line}")

            source_number, stem, options = parse_question_line(line, current_kind)
            if source_number not in answers:
                raise ValueError(
                    f"Block {block_index} question {source_number} has no answer"
                )
            answer, explanation = answers[source_number]
            block_questions.append(
                Question(
                    source_block=block_index,
                    source_number=source_number,
                    kind=current_kind,
                    stem=stem,
                    options=options,
                    answer=answer,
                    explanation=explanation,
                )
            )

        question_numbers = {question.source_number for question in block_questions}
        if question_numbers != set(answers):
            missing = sorted(set(answers) - question_numbers)
            raise ValueError(f"Block {block_index} has unmatched answers: {missing}")
        questions.extend(block_questions)

    reordered = [
        question
        for kind in TYPE_ORDER
        for question in questions
        if question.kind == kind
    ]
    for number, question in enumerate(reordered, start=1):
        question.number = number
    return reordered


def markdown_text(chapter_number: int, questions: list[Question]) -> str:
    chapter_name = CHAPTER_NAMES[chapter_number]
    lines = [f"# {chapter_name} 选择题 判断题", ""]

    for kind in TYPE_ORDER:
        typed_questions = [question for question in questions if question.kind == kind]
        if not typed_questions:
            continue
        lines.extend([TYPE_LABELS[kind], ""])
        for question in typed_questions:
            lines.append(f"**{question.number}. {question.stem}**")
            if question.options:
                lines.append(question.options)
            lines.append("")

    lines.extend(["答案与解析", ""])
    for question in questions:
        explanation = f"{question.answer}。"
        if question.explanation:
            explanation += question.explanation
        lines.extend([f"{question.number}. {explanation}", ""])

    return "\n".join(lines).rstrip() + "\n"


def set_run_font(
    run,
    *,
    size: float = 12,
    bold: bool = False,
    color: RGBColor = INK,
    italic: bool = False,
) -> None:
    run.font.name = WESTERN_FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), WESTERN_FONT)
    fonts.set(qn("w:hAnsi"), WESTERN_FONT)
    fonts.set(qn("w:eastAsia"), CHINESE_FONT)


def set_paragraph_spacing(
    paragraph,
    *,
    before: float = 0,
    after: float = 3,
    line: float = 1.18,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_bottom_rule(paragraph, color: str = "1F4E79", size: str = "16") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=GRAY)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    field_run = paragraph.add_run()
    set_run_font(field_run, size=9, color=GRAY)
    field_run._r.extend([begin, instr, separate, value, end])
    run = paragraph.add_run(" 页")
    set_run_font(run, size=9, color=GRAY)
    set_paragraph_spacing(paragraph, after=0, line=1.0)


def configure_document(doc: Document, chapter_name: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.649)
    section.right_margin = Inches(0.729)
    section.bottom_margin = Inches(0.610)
    section.left_margin = Inches(0.729)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = WESTERN_FONT
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:ascii"), WESTERN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), WESTERN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CHINESE_FONT)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.18

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(header_paragraph, after=0, line=1.0)
    run = header_paragraph.add_run(f"中国近现代史纲要 · {chapter_name}复习试卷")
    set_run_font(run, size=9, color=GRAY)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    add_page_number(footer_paragraph)


def add_title_block(
    doc: Document,
    chapter_name: str,
    questions: list[Question] | None,
    *,
    template: bool = False,
) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, after=10, line=1.0)
    title_run = title.add_run(f"中国近现代史纲要 {chapter_name}复习试卷")
    set_run_font(title_run, size=20, bold=True, color=BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=16, line=1.0)
    set_run_font(
        subtitle.add_run("选择题与判断题专项练习"),
        size=12,
        color=GRAY,
    )

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    remove_table_borders(table)
    for cell, label in zip(table.rows[0].cells, ("班级：", "姓名：", "学号：", "成绩：")):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, LIGHT_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(paragraph, after=0, line=1.0)
        set_run_font(paragraph.add_run(label), size=12)

    summary = doc.add_paragraph()
    set_paragraph_spacing(summary, before=7, after=18, line=1.15)
    if template:
        summary_text = (
            "本卷共 [总题数] 题：单选题 [数量] 题，多选题 [数量] 题，"
            "判断题 [数量] 题。请将答案写在题后空白处或另附答题纸。"
        )
    else:
        counts = {
            kind: sum(question.kind == kind for question in questions or [])
            for kind in TYPE_ORDER
        }
        summary_text = (
            f"本卷共 {len(questions or [])} 题：单选题 {counts['single']} 题，"
            f"多选题 {counts['multi']} 题，判断题 {counts['judge']} 题。"
            "请将答案写在题后空白处或另附答题纸。"
        )
    set_run_font(summary.add_run(summary_text), size=11.5, color=GRAY)

    label = doc.add_paragraph()
    set_paragraph_spacing(label, before=10, after=8, line=1.0)
    set_run_font(label.add_run("试题部分"), size=15, bold=True, color=BLUE)


def add_type_heading(doc: Document, label: str) -> None:
    heading = doc.add_paragraph()
    set_paragraph_spacing(heading, before=8, after=10, line=1.0)
    set_run_font(heading.add_run(label), size=13, bold=True, color=BLUE)
    heading.paragraph_format.keep_with_next = True
    add_bottom_rule(heading)


def add_question(doc: Document, question: Question) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=2, after=2, line=1.18)
    paragraph.paragraph_format.keep_with_next = bool(question.options)
    set_run_font(paragraph.add_run(f"{question.number}. {question.stem}"), size=12)

    if question.options:
        options = doc.add_paragraph()
        options.paragraph_format.left_indent = Inches(0.25)
        set_paragraph_spacing(options, after=4, line=1.18)
        set_run_font(options.add_run(question.options), size=12)


def add_answers(doc: Document, questions: list[Question], *, template: bool = False) -> None:
    doc.add_section(WD_SECTION.NEW_PAGE)
    heading = doc.add_paragraph()
    set_paragraph_spacing(heading, after=10, line=1.0)
    set_run_font(heading.add_run("答案与解析"), size=15, bold=True, color=BLUE)
    add_bottom_rule(heading)

    if template:
        paragraph = doc.add_paragraph()
        set_paragraph_spacing(paragraph, before=2, after=5, line=1.18)
        set_run_font(paragraph.add_run("1. [答案]。[解析]"), size=12)
        return

    for question in questions:
        paragraph = doc.add_paragraph()
        set_paragraph_spacing(paragraph, before=2, after=5, line=1.18)
        text = f"{question.number}. {question.answer}。{question.explanation}"
        set_run_font(paragraph.add_run(text), size=12)


def build_docx(chapter_number: int, questions: list[Question], output: Path) -> None:
    chapter_name = CHAPTER_NAMES[chapter_number]
    doc = Document()
    configure_document(doc, chapter_name)
    add_title_block(doc, chapter_name, questions)
    for kind in TYPE_ORDER:
        typed_questions = [question for question in questions if question.kind == kind]
        if not typed_questions:
            continue
        add_type_heading(doc, TYPE_LABELS[kind])
        for question in typed_questions:
            add_question(doc, question)
    add_answers(doc, questions)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_template(output: Path) -> None:
    doc = Document()
    configure_document(doc, "第X章")
    add_title_block(doc, "第X章", None, template=True)
    for kind in TYPE_ORDER:
        add_type_heading(doc, TYPE_LABELS[kind])
        paragraph = doc.add_paragraph()
        set_paragraph_spacing(paragraph, before=2, after=5, line=1.18)
        sample = "1. [题干]" if kind == "judge" else "1. [题干]\nA. [选项]  B. [选项]  C. [选项]  D. [选项]"
        for index, part in enumerate(sample.splitlines()):
            if index:
                paragraph.add_run().add_break()
            set_run_font(paragraph.add_run(part), size=12)
    add_answers(doc, [], template=True)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--chapter", type=int, required=True, choices=CHAPTER_NAMES)
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--template", type=Path)
    args = parser.parse_args()

    questions = parse_raw_quiz(args.input)
    output_dir = args.output_dir
    md_output = output_dir / f"chapter-{args.chapter}.md"
    docx_output = output_dir / f"chapter-{args.chapter}.docx"
    output_dir.mkdir(parents=True, exist_ok=True)
    md_output.write_text(markdown_text(args.chapter, questions), encoding="utf-8")
    build_docx(args.chapter, questions, docx_output)
    if args.template:
        build_template(args.template)

    counts = {
        kind: sum(question.kind == kind for question in questions)
        for kind in TYPE_ORDER
    }
    print(
        f"chapter={args.chapter} total={len(questions)} "
        f"single={counts['single']} multi={counts['multi']} judge={counts['judge']}"
    )
    print(md_output)
    print(docx_output)


if __name__ == "__main__":
    main()
